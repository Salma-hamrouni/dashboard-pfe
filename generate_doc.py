"""
Generate Dashboard Generator API - Documentation PFE .docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.5)

# ─── Helpers ─────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x1F, 0x49, 0x7D)
LBLUE  = RGBColor(0x2E, 0x74, 0xB5)
GREEN  = RGBColor(0x10, 0x7C, 0x10)
GRAY   = RGBColor(0xF2, 0xF2, 0xF2)
DARK   = RGBColor(0x26, 0x26, 0x26)
CODE_BG = RGBColor(0xF4, 0xF4, 0xF4)

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE
        run.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = LBLUE
        run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = GREEN
        run.font.size = Pt(12)
        run.bold = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def para(text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # light background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    pPr.append(shd)
    return p

def endpoint_box(method, path, description):
    """Colored endpoint header."""
    colors = {
        'GET':    ('E8F5E9', '1B5E20'),
        'POST':   ('E3F2FD', '0D47A1'),
        'PUT':    ('FFF8E1', 'E65100'),
        'PATCH':  ('F3E5F5', '4A148C'),
        'DELETE': ('FFEBEE', 'B71C1C'),
    }
    bg, fg = colors.get(method, ('F5F5F5', '212121'))
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    # Method badge
    r1 = p.add_run(f" {method} ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16))
    # Path
    r2 = p.add_run(f"  {path}")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.name = 'Courier New'
    r2.font.color.rgb = DARK
    # Description
    if description:
        r3 = p.add_run(f"\n     {description}")
        r3.font.size = Pt(10)
        r3.italic = True
        r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        cell._tc.tcPr.append(shd)
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            cell._tc.tcPr.append(shd)
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    doc.add_paragraph()
    return table

def page_break():
    doc.add_page_break()

def separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Cm(4)
run = p.add_run("DASHBOARD GENERATOR API")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Documentation Technique Complète")
r2.font.size = Pt(18)
r2.font.color.rgb = LBLUE
r2.italic = True

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Projet de Fin d'Études (PFE)")
r3.font.size = Pt(14)
r3.font.color.rgb = RGBColor(0x44,0x44,0x44)

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run("ASP.NET Core 8  ·  Entity Framework Core  ·  MySQL  ·  JWT  ·  Gemini AI  ·  Docker")
r4.font.size = Pt(11)
r4.font.color.rgb = RGBColor(0x66,0x66,0x66)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("2025 – 2026")
r5.font.size = Pt(12)
r5.font.color.rgb = RGBColor(0x88,0x88,0x88)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. Introduction")
separator()

h2("1.1 Présentation du Projet")
para(
    "Le Dashboard Generator est une plateforme web complète qui permet à ses utilisateurs de créer, "
    "personnaliser et partager des tableaux de bord interactifs. L'utilisateur peut connecter ses "
    "propres sources de données (fichiers CSV, bases de données SQL, APIs REST), créer des graphiques "
    "et des indicateurs clés (KPIs), et bénéficier de recommandations automatiques grâce à l'intelligence "
    "artificielle."
)

h2("1.2 Rôle du Backend")
para(
    "Ce document décrit uniquement la couche backend de l'application : l'API REST développée avec "
    "ASP.NET Core 8. Cette API est le cerveau de l'application. Elle reçoit toutes les requêtes du "
    "frontend, les traite, les sécurise, accède à la base de données, appelle l'IA si besoin, et "
    "retourne des réponses JSON structurées."
)

h2("1.3 Technologies Utilisées")
add_table(
    ["Technologie", "Rôle dans le projet"],
    [
        ["ASP.NET Core 8",      "Framework principal pour créer l'API REST"],
        ["Entity Framework Core 9", "ORM : communication avec la base de données MySQL sans écrire de SQL manuel"],
        ["MySQL 8.0",           "Base de données relationnelle principale"],
        ["JWT (JSON Web Token)","Authentification sécurisée : chaque utilisateur reçoit un token après connexion"],
        ["BCrypt",              "Hachage des mots de passe (le mot de passe réel n'est jamais stocké)"],
        ["Swagger / OpenAPI",   "Documentation interactive de l'API accessible depuis un navigateur"],
        ["Gemini AI (Google)",  "Intelligence artificielle pour recommander des graphiques et analyser les données"],
        ["Serilog",             "Journalisation structurée (logs console + fichiers tournants quotidiens)"],
        ["Docker",              "Conteneurisation : déployer l'API + MySQL en une seule commande"],
        ["xUnit",               "Framework de tests unitaires (32 tests automatisés)"],
        ["IMemoryCache",        "Cache mémoire pour éviter des requêtes répétées à la DB ou à l'IA"],
    ],
    [5, 12]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. Architecture Générale du Backend")
separator()

h2("2.1 Flux de Traitement d'une Requête")
para(
    "Chaque appel au backend suit le même chemin. Voici comment une requête voyage de "
    "l'utilisateur jusqu'à la réponse :"
)
code_block(
    "Navigateur / Application Frontend\n"
    "        │\n"
    "        ▼  HTTP Request (JSON)\n"
    "┌─────────────────────────────────┐\n"
    "│       Middleware Pipeline       │\n"
    "│  GlobalException → Security     │\n"
    "│  Headers → Serilog → CORS →    │\n"
    "│  RateLimit → Auth → AuthZ      │\n"
    "└────────────┬────────────────────┘\n"
    "             │\n"
    "             ▼\n"
    "┌─────────────────────────────────┐\n"
    "│          Controllers            │\n"
    "│  Auth / Dashboard / Widget /   │\n"
    "│  DataSource / AI / Export...   │\n"
    "└────────────┬────────────────────┘\n"
    "             │\n"
    "             ▼\n"
    "┌─────────────────────────────────┐\n"
    "│           Services              │\n"
    "│  AuthService / WidgetDataSvc   │\n"
    "│  CsvParser / SqlConnector /    │\n"
    "│  AuditService / AiService...   │\n"
    "└────────────┬────────────────────┘\n"
    "             │\n"
    "        ┌────┴────┐\n"
    "        ▼         ▼\n"
    "    MySQL DB   Gemini AI\n"
    "        │         │\n"
    "        └────┬────┘\n"
    "             ▼\n"
    "    ApiResponse<T> JSON\n"
    "        │\n"
    "        ▼\n"
    "  Frontend / Chart.js"
)

h2("2.2 Composants Principaux")

h3("Controllers")
para(
    "Un Controller est le point d'entrée d'une requête HTTP. Chaque module a son propre controller "
    "(AuthController, DashboardController, WidgetController...). Le controller reçoit la requête, "
    "vérifie les droits, appelle les services nécessaires et retourne la réponse."
)

h3("Services")
para(
    "Un Service contient la logique métier. Par exemple, AuthService gère la vérification du mot "
    "de passe et la génération du token JWT. WidgetDataService calcule les agrégations (SUM, AVG, "
    "COUNT) sur les données. Les services sont injectés dans les controllers via l'injection de "
    "dépendances (DI) d'ASP.NET Core."
)

h3("Entity Framework Core (EF Core)")
para(
    "EF Core est le pont entre le code C# et la base de données MySQL. Au lieu d'écrire des requêtes "
    "SQL manuellement, on écrit du code C# (LINQ) et EF Core traduit automatiquement en SQL. Par "
    "exemple : context.Users.Where(u => u.Email == email).FirstOrDefaultAsync() devient "
    "SELECT * FROM Users WHERE Email = '...' LIMIT 1."
)

h3("DTOs (Data Transfer Objects)")
para(
    "Un DTO est un objet qui définit exactement les données envoyées ou reçues par l'API. Il permet "
    "de ne jamais exposer directement le modèle de base de données (qui peut contenir des champs "
    "sensibles comme le mot de passe haché)."
)

h3("Middlewares")
para(
    "Un middleware est du code qui s'exécute avant chaque requête. L'API utilise :"
)
bullet("GlobalExceptionMiddleware : intercepte toutes les erreurs non gérées et retourne un JSON propre au lieu d'un crash")
bullet("SecurityHeadersMiddleware : ajoute des en-têtes de sécurité HTTP (CSP, HSTS, X-Frame-Options...)")
bullet("Serilog Request Logging : enregistre automatiquement chaque requête HTTP avec sa durée")

h3("Cache Mémoire (IMemoryCache)")
para(
    "Le cache permet de stocker temporairement des résultats coûteux en mémoire. Si le même calcul "
    "est demandé 100 fois en 5 minutes, il n'est effectué qu'une seule fois. Utilisé pour les "
    "statistiques (5 min) et les réponses IA (1 heure)."
)

h3("Audit Logs")
para(
    "Chaque action importante (connexion, création de dashboard, suppression...) est enregistrée dans "
    "la table AuditLogs avec l'ID de l'utilisateur, l'action, l'IP et la date. Permet à l'administrateur "
    "de surveiller toute l'activité de la plateforme."
)

h2("2.3 Format de Réponse Standard")
para(
    "Toutes les réponses de l'API suivent le même format enveloppe ApiResponse<T> :"
)
code_block(
    '{\n'
    '  "success": true,\n'
    '  "data": { ... },        // la vraie donnée\n'
    '  "meta": { ... },        // infos supplémentaires (source cache, pagination...)\n'
    '  "errors": null          // null si succès, liste des erreurs sinon\n'
    '}'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. MODULE AUTH
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Module Auth — Authentification")
separator()
para(
    "Ce module gère toute la sécurité d'accès à l'application : inscription, connexion, "
    "identification de l'utilisateur connecté et gestion des rôles (User / Admin)."
)

# Register
endpoint_box("POST", "/api/auth/register", "Créer un nouveau compte utilisateur")
add_table(
    ["Paramètre", "Type", "Obligatoire", "Description"],
    [
        ["email",    "string", "Oui", "Adresse email unique de l'utilisateur"],
        ["password", "string", "Oui", "Mot de passe (min. 6 caractères)"],
        ["fullName", "string", "Oui", "Nom complet affiché dans l'interface"],
        ["role",     "string", "Non", "Rôle : User (défaut) ou Admin"],
    ],
    [3, 2.5, 2.5, 8]
)
para("Fonctionnement interne :", bold=True)
bullet("1. Vérifie que l'email n'existe pas déjà dans la base de données")
bullet("2. Hache le mot de passe avec BCrypt (algorithme sécurisé — irréversible)")
bullet("3. Crée l'utilisateur en base avec UserId, CreatedAt automatiques")
bullet("4. Enregistre un AuditLog 'USER_REGISTERED' avec l'IP de l'appelant")
para("Exemple de requête :", bold=True)
code_block(
    'POST /api/auth/register\n'
    'Content-Type: application/json\n\n'
    '{\n'
    '  "email":    "alice@example.com",\n'
    '  "password": "MonMotDePasse123",\n'
    '  "fullName": "Alice Dupont"\n'
    '}'
)
para("Exemple de réponse (201 Created) :", bold=True)
code_block(
    '{\n'
    '  "success": true,\n'
    '  "data": {\n'
    '    "id":        42,\n'
    '    "email":     "alice@example.com",\n'
    '    "fullName":  "Alice Dupont",\n'
    '    "role":      "User",\n'
    '    "createdAt": "2026-05-19T10:30:00Z"\n'
    '  }\n'
    '}'
)
add_table(
    ["Code HTTP", "Signification"],
    [
        ["201 Created",      "Compte créé avec succès"],
        ["400 Bad Request",  "Email déjà utilisé ou données invalides"],
        ["429 Too Many Req", "Trop de tentatives (rate limiting : 5 req/min)"],
    ],
    [4, 13]
)
separator()

# Login
endpoint_box("POST", "/api/auth/login", "Authentifier un utilisateur et obtenir un token JWT")
para("Fonctionnement interne :", bold=True)
bullet("1. Cherche l'utilisateur par email")
bullet("2. Compare le mot de passe fourni avec le hash BCrypt stocké")
bullet("3. Si correct : génère un token JWT signé, valable 7 jours")
bullet("4. Le token contient (encodé) : ID utilisateur, email, rôle")
bullet("5. AuditLog 'USER_LOGIN' si succès, 'USER_LOGIN_FAILED' si échec")
bullet("6. Rate limiting anti brute-force : 5 tentatives max par minute par IP")
para("Exemple de réponse (200 OK) :", bold=True)
code_block(
    '{\n'
    '  "success": true,\n'
    '  "data": {\n'
    '    "token": "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",\n'
    '    "user": {\n'
    '      "id":     42,\n'
    '      "email":  "alice@example.com",\n'
    '      "role":   "User"\n'
    '    }\n'
    '  }\n'
    '}'
)
para(
    "Le frontend stocke ce token et l'envoie dans chaque requête suivante via l'en-tête HTTP : "
    "Authorization: Bearer eyJhbGciOi...",
    italic=True
)
separator()

# Me
endpoint_box("GET", "/api/auth/me", "Récupérer le profil de l'utilisateur connecté")
para(
    "Ce endpoint lit l'identité de l'utilisateur depuis son token JWT (sans requête DB supplémentaire "
    "pour l'identification) puis charge son profil complet depuis la base. Le frontend l'appelle au "
    "démarrage pour savoir qui est connecté."
)
separator()

# Users (Admin)
endpoint_box("GET", "/api/auth/users", "Lister tous les utilisateurs — Admin uniquement")
para(
    "Utilise AsNoTracking() pour une requête en lecture seule optimisée (EF Core ne surveille pas "
    "les changements sur les objets retournés, ce qui économise de la mémoire et du temps)."
)
separator()

# Change role
endpoint_box("PUT", "/api/auth/users/{id}/role", "Modifier le rôle d'un utilisateur — Admin uniquement")
para("Exemple : passer un utilisateur de 'User' à 'Admin'.")
code_block(
    'PUT /api/auth/users/42/role\n'
    'Authorization: Bearer <admin-token>\n\n'
    '{ "role": "Admin" }'
)
para(
    "Utilise ExecuteUpdateAsync : met à jour uniquement le champ Role en une seule requête SQL "
    "sans charger tout l'objet User en mémoire. Plus efficace qu'un UPDATE complet.",
    italic=True
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. MODULE DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. Module Dashboard — Tableaux de Bord")
separator()
para(
    "Le Dashboard est le conteneur principal de l'application. C'est la 'page' sur laquelle "
    "l'utilisateur place ses graphiques (widgets). Chaque utilisateur possède ses propres "
    "dashboards, invisibles des autres utilisateurs."
)

endpoint_box("GET", "/api/dashboard", "Lister les dashboards de l'utilisateur avec pagination et recherche")
para("Paramètres de requête disponibles :", bold=True)
add_table(
    ["Paramètre", "Défaut", "Description"],
    [
        ["page",     "1",          "Numéro de page (commence à 1)"],
        ["pageSize", "10",         "Nombre de dashboards par page (max 50)"],
        ["search",   "(vide)",     "Filtre par nom (recherche partielle, insensible à la casse)"],
        ["sortBy",   "createdAt",  "Colonne de tri : name, createdAt, updatedAt"],
        ["sortDesc", "true",       "Ordre décroissant (true) ou croissant (false)"],
    ],
    [3, 2.5, 11.5]
)
para(
    "Exemple : GET /api/dashboard?page=2&pageSize=5&search=ventes&sortBy=name retourne "
    "la page 2 des dashboards dont le nom contient 'ventes', triés alphabétiquement.",
    italic=True
)
para("Exemple de réponse :", bold=True)
code_block(
    '{\n'
    '  "success": true,\n'
    '  "data": [\n'
    '    { "id": 1, "name": "Ventes 2026", "isPublic": false, "widgetCount": 5 },\n'
    '    { "id": 2, "name": "Ventes Région Sud", "isPublic": true, "widgetCount": 3 }\n'
    '  ],\n'
    '  "meta": {\n'
    '    "page": 2, "pageSize": 5,\n'
    '    "totalCount": 12, "totalPages": 3\n'
    '  }\n'
    '}'
)
separator()

endpoint_box("POST", "/api/dashboard", "Créer un nouveau tableau de bord")
code_block(
    'POST /api/dashboard\n'
    'Authorization: Bearer <token>\n\n'
    '{\n'
    '  "name":        "Tableau de bord Ventes Q2",\n'
    '  "description": "Suivi des ventes du 2ème trimestre",\n'
    '  "isPublic":    false\n'
    '}'
)
separator()

endpoint_box("GET", "/api/dashboard/{id}", "Récupérer un dashboard avec tous ses widgets")
para(
    "Utilise Include(d => d.Widgets) : EF Core génère automatiquement une jointure SQL pour "
    "charger le dashboard et ses widgets en une seule requête. Vérifie que l'utilisateur est "
    "bien le propriétaire avant de retourner les données (sécurité ownership)."
)
separator()

endpoint_box("PUT", "/api/dashboard/{id}", "Modifier le nom, la description ou la visibilité")
endpoint_box("DELETE", "/api/dashboard/{id}", "Supprimer un dashboard et tous ses widgets (cascade automatique)")
separator()

endpoint_box("PATCH", "/api/dashboard/{id}/toggle-public", "Basculer la visibilité public/privé en un seul appel")
para(
    "Inverse simplement la valeur booléenne IsPublic. Si le dashboard était privé (false), "
    "il devient public (true) et vice versa. Un dashboard public peut être accédé sans "
    "authentification via son lien de partage."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. MODULE WIDGET + DATA ENGINE
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. Module Widget — Graphiques et Composants")
separator()
para(
    "Un widget est un composant graphique placé sur un dashboard. Il peut être un graphique "
    "en barres, en courbes, un camembert, un tableau, ou un KPI (indicateur clé). "
    "Chaque widget est lié à une source de données et possède une configuration qui définit "
    "comment afficher les données."
)

endpoint_box("GET", "/api/widget/dashboard/{dashboardId}", "Récupérer tous les widgets d'un dashboard")
endpoint_box("POST", "/api/widget", "Créer un nouveau widget")
code_block(
    'POST /api/widget\n\n'
    '{\n'
    '  "dashboardId":  1,\n'
    '  "title":        "Ventes par Région",\n'
    '  "type":         "bar",\n'
    '  "config": {\n'
    '    "dataSourceId": 3,\n'
    '    "xAxis":        "region",\n'
    '    "yAxis":        "montant_vente",\n'
    '    "aggregation":  "SUM"\n'
    '  },\n'
    '  "positionX": 0, "positionY": 0,\n'
    '  "width": 6,     "height": 4\n'
    '}'
)
para("Types de widgets supportés :", bold=True)
add_table(
    ["Type",  "Description"],
    [
        ["bar",     "Graphique en barres verticales ou horizontales"],
        ["line",    "Graphique en courbes (idéal pour les séries temporelles)"],
        ["pie",     "Graphique en camembert (répartition en pourcentages)"],
        ["scatter", "Nuage de points (corrélation entre deux variables)"],
        ["kpi",     "Indicateur clé : affiche une valeur unique avec label (ex: 2.3M €)"],
        ["table",   "Tableau de données paginé"],
    ],
    [3, 14]
)
separator()

endpoint_box("PUT", "/api/widget/{id}", "Modifier le titre, la taille ou la position du widget")
endpoint_box("PATCH", "/api/widget/{id}/config", "Mettre à jour uniquement la configuration JSON (axes, couleurs, source)")
endpoint_box("DELETE", "/api/widget/{id}", "Supprimer un widget du dashboard")
separator()

h2("5.1 Widget Data Engine — GET /api/widget/{id}/data")
para(
    "C'est l'endpoint le plus important et le plus complexe de toute l'API. Il transforme "
    "des données brutes (CSV, SQL, REST) en un JSON directement utilisable par Chart.js "
    "pour afficher un graphique sur le frontend."
)

h3("Pipeline de Traitement Complet")
code_block(
    "┌─────────────────────────────────────────┐\n"
    "│  Données Brutes (CSV / SQL / REST)      │\n"
    "│  Stockées en JSON dans la base de données│\n"
    "└─────────────────┬───────────────────────┘\n"
    "                  │\n"
    "                  ▼\n"
    "┌─────────────────────────────────────────┐\n"
    "│  Lecture Config Widget                  │\n"
    "│  xAxis: 'region'                        │\n"
    "│  yAxis: 'montant_vente'                 │\n"
    "│  aggregation: 'SUM'                     │\n"
    "│  filters: [{ col: 'annee', val: '2026' }]│\n"
    "└─────────────────┬───────────────────────┘\n"
    "                  │\n"
    "                  ▼\n"
    "┌─────────────────────────────────────────┐\n"
    "│  Étape 1 : FILTRAGE                     │\n"
    "│  Garde uniquement les lignes où         │\n"
    "│  annee == '2026'                        │\n"
    "└─────────────────┬───────────────────────┘\n"
    "                  │\n"
    "                  ▼\n"
    "┌─────────────────────────────────────────┐\n"
    "│  Étape 2 : GROUPEMENT (GROUP BY)        │\n"
    "│  Groupe les lignes par valeur de        │\n"
    "│  la colonne xAxis (= 'region')          │\n"
    "│  Nord: [45000, 32000, 28000...]         │\n"
    "│  Sud:  [61000, 44000, 37000...]         │\n"
    "└─────────────────┬───────────────────────┘\n"
    "                  │\n"
    "                  ▼\n"
    "┌─────────────────────────────────────────┐\n"
    "│  Étape 3 : AGRÉGATION                   │\n"
    "│  SUM   → additionne les valeurs         │\n"
    "│  AVG   → calcule la moyenne             │\n"
    "│  COUNT → compte le nombre de lignes     │\n"
    "│  MIN   → valeur minimum du groupe       │\n"
    "│  MAX   → valeur maximum du groupe       │\n"
    "└─────────────────┬───────────────────────┘\n"
    "                  │\n"
    "                  ▼\n"
    "┌─────────────────────────────────────────┐\n"
    "│  Étape 4 : FORMAT Chart.js              │\n"
    "│  labels:   ['Nord', 'Sud', 'Est']       │\n"
    "│  datasets: [{ data: [105000, 142000,    │\n"
    "│              98000] }]                  │\n"
    "└─────────────────────────────────────────┘"
)

h3("Les 5 Types d'Agrégation")
add_table(
    ["Agrégation", "Calcul", "Exemple d'utilisation"],
    [
        ["SUM",   "Additionne toutes les valeurs du groupe",  "Chiffre d'affaires total par région"],
        ["AVG",   "Calcule la moyenne des valeurs",           "Note moyenne par catégorie de produit"],
        ["COUNT", "Compte le nombre de lignes du groupe",     "Nombre de commandes par mois"],
        ["MIN",   "Retourne la valeur la plus petite",        "Prix minimum par fournisseur"],
        ["MAX",   "Retourne la valeur la plus grande",        "Score maximum par équipe"],
        ["NONE",  "Retourne les valeurs brutes sans calcul",  "Affichage de données déjà agrégées"],
    ],
    [2.5, 7, 7.5]
)

h3("Exemple Complet : Widget KPI")
para(
    "Pour un widget de type 'kpi' qui affiche le chiffre d'affaires total, "
    "le Data Engine retourne une valeur scalaire formatée intelligemment :"
)
code_block(
    '// Si le total est 2 347 891 €\n'
    '{\n'
    '  "type":  "kpi",\n'
    '  "title": "Chiffre d\'Affaires Total",\n'
    '  "value": "2.3M",     // FormatHuman() formate les grands nombres\n'
    '  "raw":   2347891\n'
    '}'
)

h3("Exemple Complet : Widget Bar Chart")
code_block(
    '// Résultat pour un graphique en barres "Ventes par Région"\n'
    '{\n'
    '  "type": "bar",\n'
    '  "labels": ["Nord", "Sud", "Est", "Ouest"],\n'
    '  "datasets": [\n'
    '    {\n'
    '      "label": "montant_vente (SUM)",\n'
    '      "data":  [105000, 142000, 98000, 76000]\n'
    '    }\n'
    '  ]\n'
    '}'
)
para(
    "Ce JSON est directement passé à Chart.js sur le frontend sans aucune transformation. "
    "Le frontend n'a aucune logique de calcul — tout est géré côté backend.",
    italic=True
)

h3("Formatage Intelligent des Nombres (FormatHuman)")
para(
    "La fonction FormatHuman() formate les grands nombres de façon lisible, avec la culture "
    "invariante pour éviter les problèmes de localisation (virgule vs point décimal) :"
)
add_table(
    ["Valeur brute", "Résultat FormatHuman()", "Règle"],
    [
        ["2 347 891",  "2.3M",     "≥ 1 000 000 → divise par 1M, suffix M"],
        ["45 600",     "45.6K",    "≥ 1 000 → divise par 1K, suffix K"],
        ["6 000",      "6K",       "Si entier exact, pas de décimale"],
        ["342.75",     "342.75",   "< 1000 → affiche la valeur directement"],
    ],
    [4, 4, 9]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. MODULE DATASOURCE
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Module DataSource — Sources de Données")
separator()
para(
    "Une DataSource est la 'alimentation' des widgets. Elle représente une source de données "
    "réelle : un fichier CSV, une base de données SQL, ou une API REST. Les données sont "
    "récupérées, analysées et stockées en cache JSON dans la base pour être réutilisées "
    "rapidement par les widgets."
)

endpoint_box("POST", "/api/datasource/upload-csv", "Importer un fichier CSV comme source de données")
h3("5 Niveaux de Validation Sécurité")
add_table(
    ["Niveau", "Vérification", "Raison"],
    [
        ["1 - Taille",      "Max 10 MB",                            "Éviter la surcharge du serveur"],
        ["2 - Extension",   "Uniquement .csv et .tsv",              "Première barrière simple"],
        ["3 - MIME Type",   "Whitelist : text/csv, text/plain...",  "Le navigateur peut mentir sur le type"],
        ["4 - Magic Bytes", "Lit les 8 premiers octets du fichier", "Un attaquant peut renommer un .exe en .csv"],
        ["5 - Contenu",     "Max 5% de caractères non-imprimables", "Détecte les binaires partiellement déguisés"],
    ],
    [3, 6.5, 7.5]
)
para("Fichiers binaires rejetés par Magic Bytes :", bold=True)
bullet("ZIP / Office (.docx, .xlsx) : octets 50 4B")
bullet("JPEG : octets FF D8 FF")
bullet("PNG : octets 89 50 4E 47")
bullet("PDF : octets 25 50 44 46")
bullet("EXE/PE Windows : octets 4D 5A")
bullet("ELF Linux : octets 7F 45 4C 46")
para(
    "Rate limiting : maximum 5 uploads par minute par utilisateur pour éviter les abus. "
    "Taille maximale globale configurée à 20 MB via FormOptions.",
    italic=True
)
separator()

endpoint_box("POST", "/api/datasource/connect-sql", "Connecter une base de données MySQL ou SQL Server")
code_block(
    'POST /api/datasource/connect-sql\n\n'
    '{\n'
    '  "name":     "Base Ventes Production",\n'
    '  "server":   "192.168.1.10",\n'
    '  "database": "sales_db",\n'
    '  "username": "reporter",\n'
    '  "password": "...",\n'
    '  "port":     3306,\n'
    '  "query":    "SELECT region, SUM(amount) as total FROM orders GROUP BY region"\n'
    '}'
)
para(
    "Le backend teste d'abord la connexion, puis exécute la requête SQL et stocke les résultats "
    "en JSON dans la colonne CachedDataJson de la base.",
    italic=True
)
separator()

endpoint_box("POST", "/api/datasource/connect-rest", "Connecter une API REST externe comme source")
code_block(
    'POST /api/datasource/connect-rest\n\n'
    '{\n'
    '  "name":     "API Météo Paris",\n'
    '  "endpoint": "https://api.example.com/weather",\n'
    '  "method":   "GET",\n'
    '  "headers":  { "Authorization": "Bearer xxx" },\n'
    '  "dataPath": "data.results"  // JSON path pour extraire le tableau\n'
    '}'
)
separator()

endpoint_box("POST", "/api/datasource/{id}/refresh", "Re-synchroniser les données depuis la source originale")
para(
    "Selon le type de source, relance le bon connecteur : re-parse le fichier CSV, "
    "re-exécute la requête SQL, ou rappelle l'API REST. Met à jour CachedDataJson et "
    "LastRefreshedAt en base."
)
para(
    "Refresh automatique : un service en arrière-plan (DataSourceRefreshService) "
    "re-synchronise automatiquement toutes les sources toutes les 30 minutes (configurable).",
    italic=True
)
separator()

endpoint_box("POST", "/api/datasource/{id}/preview", "Afficher un aperçu des 5 premières lignes de données")
endpoint_box("GET", "/api/datasource", "Lister toutes les sources de données de l'utilisateur")
endpoint_box("GET", "/api/datasource/{id}", "Récupérer le détail d'une source de données")
endpoint_box("DELETE", "/api/datasource/{id}", "Supprimer une source (supprime aussi le fichier CSV physique si CSV)")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. MODULE DATASET
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Module Dataset — Analyse de Données")
separator()
para(
    "Le module Dataset est orienté exploration et analyse de données. Contrairement à "
    "DataSource (qui alimente les widgets), Dataset permet d'uploader des fichiers pour "
    "les analyser, obtenir des statistiques descriptives et des recommandations."
)

endpoint_box("POST", "/api/datasets/upload", "Uploader un fichier de données pour analyse")
para(
    "Supporte CSV, JSON et Excel. Après upload, l'API analyse automatiquement les colonnes "
    "(type numérique, texte, date), calcule des statistiques de base et retourne un profil "
    "complet du dataset."
)

endpoint_box("POST", "/api/datasets/analyze", "Analyser statistiquement un dataset uploadé")
para("Retourne pour chaque colonne :")
bullet("Type détecté (numérique, catégoriel, date, texte)")
bullet("Nombre de valeurs nulles / manquantes")
bullet("Valeurs min, max, moyenne, médiane (pour les colonnes numériques)")
bullet("Top 5 valeurs les plus fréquentes (pour les colonnes catégorielles)")

endpoint_box("POST", "/api/datasets/recommend", "Recommander les meilleurs graphiques pour le dataset")
para(
    "Analyse la structure du dataset et suggère automatiquement quels graphiques sont "
    "les plus adaptés : bar chart pour les catégories, line chart pour les dates, "
    "scatter pour les corrélations, etc."
)

endpoint_box("POST", "/api/datasets/ai-analyze", "Analyser le dataset avec l'IA Gemini")
para(
    "Envoie un résumé du dataset à l'IA Gemini pour obtenir une analyse en langage naturel : "
    "tendances, anomalies, insights métier et recommandations d'actions."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. MODULE AI
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Module AI — Intelligence Artificielle")
separator()
para(
    "Ce module intègre Gemini AI (Google) pour apporter une intelligence automatique "
    "à la plateforme. L'IA analyse les données de l'utilisateur et propose des "
    "recommandations pertinentes sans que l'utilisateur n'ait à configurer quoi que ce soit manuellement."
)

h2("8.1 Architecture de l'Intégration IA")
code_block(
    "Utilisateur envoie ses colonnes de données\n"
    "        │\n"
    "        ▼\n"
    "AiController reçoit la requête\n"
    "        │\n"
    "        ▼\n"
    "Vérifie le cache IMemoryCache\n"
    "(clé = SHA-256 des colonnes + endpoint)\n"
    "        │\n"
    "   ┌────┴────┐\n"
    "   │         │\n"
    " Cache HIT  Cache MISS\n"
    "   │         │\n"
    "   │         ▼\n"
    "   │   Construit le prompt Gemini\n"
    "   │         │\n"
    "   │         ▼\n"
    "   │   HTTP POST → Gemini API\n"
    "   │   (gemini-2.5-flash)\n"
    "   │         │\n"
    "   │         ▼\n"
    "   │   Parse la réponse JSON\n"
    "   │         │\n"
    "   │         ▼\n"
    "   │   Stocke en cache (1h TTL)\n"
    "   │         │\n"
    "   └────┬────┘\n"
    "        │\n"
    "        ▼\n"
    "   Retourne ApiResponse<JsonElement>"
)

h2("8.2 Cache Intelligent")
para(
    "Les appels à Gemini coûtent de l'argent et prennent du temps (latence réseau). "
    "Pour les colonnes identiques, la recommandation sera toujours la même. "
    "Le cache évite donc des appels inutiles :"
)
bullet("Clé de cache : SHA-256 des noms et types de colonnes + nom de l'endpoint")
bullet("Si les mêmes colonnes sont soumises 50 fois dans l'heure, Gemini n'est appelé qu'une fois")
bullet("TTL : 1 heure (AbsoluteExpiration) + 20 minutes de glissement (SlidingExpiration)")
bullet("Rate limiting : 15 requêtes/minute maximum (toutes IA confondues)")
separator()

endpoint_box("POST", "/api/ai/recommend-chart", "Recommander le type de graphique le plus adapté")
code_block(
    'POST /api/ai/recommend-chart\n\n'
    '{\n'
    '  "columns": [\n'
    '    { "name": "date_commande", "type": "date" },\n'
    '    { "name": "montant",       "type": "number" },\n'
    '    { "name": "region",        "type": "string" }\n'
    '  ]\n'
    '}'
)
para("Exemple de réponse IA :", bold=True)
code_block(
    '{\n'
    '  "success": true,\n'
    '  "data": {\n'
    '    "recommendedChart": "line",\n'
    '    "reason": "La présence d\'une colonne date et d\'une colonne numérique indique\n'
    '               une série temporelle, idéale pour un graphique en courbes.",\n'
    '    "alternatives": ["bar", "scatter"],\n'
    '    "xAxis": "date_commande",\n'
    '    "yAxis": "montant"\n'
    '  },\n'
    '  "meta": { "source": "cache" }\n'
    '}'
)
separator()

endpoint_box("POST", "/api/ai/suggest-kpis", "Proposer des indicateurs clés de performance")
para("L'IA analyse les colonnes et propose des KPIs métier pertinents à afficher sur le dashboard.")
code_block(
    '// Réponse pour un dataset de ventes\n'
    '{\n'
    '  "kpis": [\n'
    '    {\n'
    '      "title":       "Chiffre d\'Affaires Total",\n'
    '      "column":      "montant",\n'
    '      "aggregation": "sum",\n'
    '      "format":      "currency",\n'
    '      "description": "Somme de tous les montants de vente"\n'
    '    },\n'
    '    {\n'
    '      "title":       "Commandes Moyennes",\n'
    '      "column":      "montant",\n'
    '      "aggregation": "avg",\n'
    '      "format":      "number",\n'
    '      "description": "Montant moyen par commande"\n'
    '    }\n'
    '  ]\n'
    '}'
)
separator()

endpoint_box("POST", "/api/ai/analyze", "Analyser les données et détecter tendances et anomalies")
para("Envoie les colonnes ET un aperçu des données (10 premières lignes) à Gemini pour une analyse contextuelle.")
code_block(
    '// Réponse exemple\n'
    '{\n'
    '  "summary": "Dataset de 1247 commandes sur 12 mois avec 3 régions.",\n'
    '  "insights": [\n'
    '    "La région Nord représente 42% du chiffre d\'affaires total.",\n'
    '    "Le mois de décembre affiche systématiquement +35% vs la moyenne."\n'
    '  ],\n'
    '  "trends":    ["Croissance régulière de 8% trimestre après trimestre"],\n'
    '  "anomalies": ["Pic inhabituel le 15 mars (+520% vs moyenne journalière)"],\n'
    '  "recommendations": ["Renforcer les stocks en novembre pour anticiper décembre"]\n'
    '}'
)
separator()

endpoint_box("POST", "/api/ai/chat", "Décrire en langage naturel ce qu'on veut visualiser")
para(
    "L'utilisateur décrit en français ce qu'il veut voir. L'IA génère automatiquement "
    "la configuration complète du widget."
)
code_block(
    'POST /api/ai/chat\n\n'
    '{\n'
    '  "message": "Je veux voir les ventes par région en camembert",\n'
    '  "columns": [\n'
    '    { "name": "region",  "type": "string" },\n'
    '    { "name": "montant", "type": "number" }\n'
    '  ]\n'
    '}'
)
code_block(
    '// Réponse : config widget prête à l\'emploi\n'
    '{\n'
    '  "type":        "pie",\n'
    '  "title":       "Ventes par Région",\n'
    '  "xAxis":       "region",\n'
    '  "yAxis":       "montant",\n'
    '  "aggregation": "sum",\n'
    '  "filters":     [],\n'
    '  "explanation": "Camembert avec SUM des montants groupés par région"\n'
    '}'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. MODULE EXPORT
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. Module Export — Sauvegarde et Partage")
separator()
para(
    "Ce module permet d'exporter des dashboards complets pour les sauvegarder, "
    "les partager ou les importer sur un autre compte."
)

endpoint_box("GET", "/api/export/{dashboardId}/json", "Exporter un dashboard complet au format JSON")
para(
    "Retourne le dashboard avec tous ses widgets sérialisés en JSON indenté. "
    "Ce fichier peut être sauvegardé puis réimporté plus tard. "
    "Rate limiting : 20 exports/minute."
)
code_block(
    '// Extrait du JSON exporté\n'
    '{\n'
    '  "id": 1,\n'
    '  "name": "Tableau de bord Ventes",\n'
    '  "createdAt": "2026-05-10T...",\n'
    '  "widgets": [\n'
    '    {\n'
    '      "title": "Ventes par Région",\n'
    '      "type":  "bar",\n'
    '      "config": { "xAxis": "region", "yAxis": "montant", "aggregation": "SUM" }\n'
    '    }\n'
    '  ]\n'
    '}'
)
separator()

endpoint_box("GET", "/api/export/{dashboardId}/pdf", "Générer un rapport PDF du dashboard")
para(
    "Génère un fichier PDF avec les informations du dashboard et la liste structurée "
    "de tous ses widgets. Rate limiting : 20 exports/minute."
)
separator()

endpoint_box("POST", "/api/export/import", "Importer un dashboard depuis un fichier JSON exporté")
para(
    "Parse le JSON fourni, recrée le dashboard et tous ses widgets dans la base de données, "
    "en les associant à l'utilisateur connecté (même si le JSON provenait d'un autre compte). "
    "Fonctionnalité de template : partager un dashboard modèle entre utilisateurs."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. MODULE DASHBOARDSHARE
# ═══════════════════════════════════════════════════════════════════════════════
h1("10. Module DashboardShare — Partage par Lien")
separator()
para(
    "Ce module permet de partager un dashboard avec des personnes extérieures via un "
    "lien unique, sans qu'elles aient besoin de créer un compte. On peut définir "
    "les permissions (lecture seule ou édition) et une date d'expiration."
)

h2("10.1 Fonctionnement du Partage")
code_block(
    "1. Propriétaire génère un lien\n"
    "   POST /api/dashboard-share\n"
    "   → Token GUID unique généré : 'a3f8b2c1-...'\n"
    "   → Lien : https://app.com/share/a3f8b2c1-...\n"
    "\n"
    "2. Destinataire ouvre le lien\n"
    "   GET /api/dashboard-share/access/a3f8b2c1-...\n"
    "   → Backend vérifie : token existe ? non expiré ?\n"
    "   → Si OK : retourne le dashboard selon les permissions\n"
    "\n"
    "3. Token expire automatiquement à la date fixée\n"
    "   → Toute tentative d'accès retourne 404"
)

endpoint_box("POST", "/api/dashboard-share", "Créer un lien de partage sécurisé")
code_block(
    'POST /api/dashboard-share\n\n'
    '{\n'
    '  "dashboardId": 1,\n'
    '  "permission":  "view",    // "view" (lecture) ou "edit" (édition)\n'
    '  "expiresAt":   "2026-06-01T00:00:00Z"  // null = jamais\n'
    '}'
)
code_block(
    '// Réponse\n'
    '{\n'
    '  "id":          5,\n'
    '  "token":       "a3f8b2c1-9d4e-4f2a-b8c7-1234567890ab",\n'
    '  "permission":  "view",\n'
    '  "expiresAt":   "2026-06-01T00:00:00Z",\n'
    '  "shareUrl":    "https://app.com/share/a3f8b2c1-9d4e-4f2a-b8c7-1234567890ab"\n'
    '}'
)
separator()

endpoint_box("GET", "/api/dashboard-share/access/{token}", "Accéder à un dashboard via lien (sans compte requis)")
para("Aucune authentification JWT requise. Le token est la clé d'accès.")
add_table(
    ["Vérification", "Si échec"],
    [
        ["Le token existe en base",             "404 Not Found"],
        ["ExpiresAt est null ou dans le futur", "404 Not Found (lien expiré)"],
        ["Permission 'view' ou 'edit'",         "Données retournées selon permission"],
    ],
    [8, 9]
)
separator()

endpoint_box("GET", "/api/dashboard-share/{dashboardId}", "Lister tous les liens de partage actifs d'un dashboard")
endpoint_box("DELETE", "/api/dashboard-share/{id}", "Révoquer un lien de partage (le rend immédiatement inaccessible)")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. MODULE DASHBOARDVERSION
# ═══════════════════════════════════════════════════════════════════════════════
h1("11. Module DashboardVersion — Historique et Restauration")
separator()
para(
    "Ce module est analogue à un système de contrôle de version simplifié (comme Git, "
    "mais pour les dashboards). Il permet de sauvegarder l'état complet d'un dashboard "
    "à un instant donné, et de le restaurer à tout moment — comme un Ctrl+Z illimité."
)

h2("11.1 Analogie Git")
add_table(
    ["Concept Git", "Équivalent Dashboard Version"],
    [
        ["git commit -m 'message'",    "POST /save → sauvegarde une version avec un label"],
        ["git log",                     "GET /{dashboardId} → liste toutes les versions"],
        ["git show v3",                 "GET /{dashboardId}/3 → détail d'une version spécifique"],
        ["git checkout v2",             "POST /restore → restaure le dashboard à une version"],
        ["git tag -d",                  "DELETE → supprime une version de l'historique"],
    ],
    [5, 12]
)

endpoint_box("POST", "/api/dashboard-version/{dashboardId}/save", "Sauvegarder l'état actuel du dashboard")
code_block(
    'POST /api/dashboard-version/1/save\n\n'
    '{ "label": "Avant la présentation client du 20 mai" }'
)
para("Ce que le backend fait :", bold=True)
bullet("1. Charge le dashboard avec tous ses widgets")
bullet("2. Sérialise l'ensemble en JSON (snapshot complète)")
bullet("3. Stocke cette snapshot en base avec un numéro de version auto-incrémenté")
code_block(
    '// Réponse\n'
    '{\n'
    '  "id":          15,\n'
    '  "dashboardId": 1,\n'
    '  "version":     4,\n'
    '  "label":       "Avant la présentation client du 20 mai",\n'
    '  "createdAt":   "2026-05-20T09:00:00Z"\n'
    '}'
)
separator()

endpoint_box("GET", "/api/dashboard-version/{dashboardId}", "Lister tout l'historique des versions")
code_block(
    '// Réponse\n'
    '[\n'
    '  { "version": 4, "label": "Avant présentation client",    "createdAt": "2026-05-20..." },\n'
    '  { "version": 3, "label": "Après révision des KPIs",      "createdAt": "2026-05-18..." },\n'
    '  { "version": 2, "label": "Version initiale avec 3 charts","createdAt": "2026-05-15..." },\n'
    '  { "version": 1, "label": "Première sauvegarde",          "createdAt": "2026-05-10..." }\n'
    ']'
)
separator()

endpoint_box("POST", "/api/dashboard-version/{dashboardId}/{versionNumber}/restore", "Restaurer une version")
para("C'est l'opération la plus puissante. Le backend effectue :", bold=True)
bullet("1. Charge la snapshot JSON de la version demandée")
bullet("2. Supprime tous les widgets actuels du dashboard")
bullet("3. Recrée chaque widget depuis la snapshot (avec les mêmes configurations)")
bullet("4. Met à jour UpdatedAt du dashboard")
para(
    "Cas d'usage : l'utilisateur a supprimé accidentellement plusieurs widgets importants. "
    "En 1 clic, il restaure la version sauvegardée avant la suppression.",
    italic=True
)
separator()

endpoint_box("GET", "/api/dashboard-version/{dashboardId}/{versionNumber}", "Voir le contenu d'une version spécifique")
endpoint_box("DELETE", "/api/dashboard-version/{id}", "Supprimer une version de l'historique pour libérer de l'espace")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. MODULE STATS
# ═══════════════════════════════════════════════════════════════════════════════
h1("12. Module Stats — Statistiques Système (Admin)")
separator()
para(
    "Ce module est réservé aux administrateurs. Il fournit une vue complète de l'état "
    "et de l'activité de la plateforme : nombre d'utilisateurs, de dashboards, activité "
    "récente, sources de données obsolètes, et logs d'audit."
)

endpoint_box("GET", "/api/stats/overview", "Vue d'ensemble de toute la plateforme — mis en cache 5 minutes")
para(
    "C'est le tableau de bord de l'administrateur. Il est mis en cache 5 minutes : "
    "si 100 admins l'appellent en 5 minutes, la base de données n'est interrogée qu'une fois."
)
code_block(
    '{\n'
    '  "totals": {\n'
    '    "users":       142,\n'
    '    "dashboards":   89,\n'
    '    "widgets":     512,\n'
    '    "datasets":     34,\n'
    '    "dataSources":  67\n'
    '  },\n'
    '  "activity": {\n'
    '    "newUsersLast30Days":         23,\n'
    '    "dashboardsCreatedLast7Days":  8,\n'
    '    "activeShares":               15,\n'
    '    "publicDashboards":           12\n'
    '  },\n'
    '  "generatedAt": "2026-05-19T10:30:00Z",\n'
    '  "meta": { "source": "cache" }\n'
    '}'
)
separator()

endpoint_box("GET", "/api/stats/users", "Répartition des utilisateurs par rôle + 10 derniers inscrits")
endpoint_box("GET", "/api/stats/dashboards", "Top 10 dashboards par nombre de widgets + versions récentes")
endpoint_box("GET", "/api/stats/datasources", "Répartition par type + sources non rafraîchies depuis 24h")

endpoint_box("GET", "/api/stats/audit", "Activité d'audit de la dernière semaine")
para("Permet à l'admin de surveiller toute l'activité et détecter les comportements suspects.")
code_block(
    '{\n'
    '  "totalActionsLast7Days": 1847,\n'
    '  "topActions": [\n'
    '    { "action": "WIDGET_CREATED",    "count": 312 },\n'
    '    { "action": "DASHBOARD_UPDATED", "count": 198 },\n'
    '    { "action": "USER_LOGIN",        "count": 156 }\n'
    '  ],\n'
    '  "mostActiveUsers": [\n'
    '    { "userId": 5, "actions": 89 },\n'
    '    { "userId": 12, "actions": 67 }\n'
    '  ],\n'
    '  "recentLogs": [\n'
    '    {\n'
    '      "userId": 5, "action": "WIDGET_DELETED",\n'
    '      "entityType": "Widget", "entityId": "42",\n'
    '      "ipAddress": "192.168.1.105",\n'
    '      "createdAt": "2026-05-19T10:28:00Z"\n'
    '    }\n'
    '  ]\n'
    '}'
)
separator()

endpoint_box("GET", "/api/stats/widgets", "Popularité des types de widgets sur toute la plateforme")
code_block(
    '{\n'
    '  "byType": [\n'
    '    { "type": "bar",   "count": 234 },\n'
    '    { "type": "line",  "count": 156 },\n'
    '    { "type": "kpi",   "count": 89  },\n'
    '    { "type": "pie",   "count": 78  },\n'
    '    { "type": "table", "count": 45  }\n'
    '  ]\n'
    '}'
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. MODULE HEALTH
# ═══════════════════════════════════════════════════════════════════════════════
h1("13. Module Health — Surveillance de l'API")
separator()
para(
    "Ces endpoints sont utilisés par les systèmes de monitoring automatique (Docker healthcheck, "
    "Kubernetes probes, Nginx upstreams). Ils permettent de vérifier en temps réel si l'API "
    "fonctionne correctement et si toutes ses dépendances sont disponibles."
)

endpoint_box("GET", "/api/health", "Liveness probe — L'API répond-elle ? (Aucune authentification)")
para(
    "C'est la question la plus simple : 'L'API est-elle en vie ?' "
    "Toujours HTTP 200 si le serveur tourne. Docker et Kubernetes l'utilisent pour "
    "décider s'il faut redémarrer le container."
)
code_block(
    '{ "status": "alive", "timestamp": "2026-05-19T10:30:00Z", "version": "1.0" }'
)
separator()

endpoint_box("GET", "/api/health/ready", "Readiness probe — Toutes les dépendances sont-elles prêtes ? (Aucune auth)")
para(
    "Va plus loin que liveness. Teste si la base de données est accessible. "
    "Si la DB est down, retourne HTTP 503 pour indiquer au load balancer de "
    "ne plus envoyer de trafic vers cette instance."
)
code_block(
    '// HTTP 200 - Tout va bien\n'
    '{\n'
    '  "status": "ready",\n'
    '  "checks": {\n'
    '    "database": { "ok": true, "latencyMs": 3 }\n'
    '  }\n'
    '}\n\n'
    '// HTTP 503 - Base de données inaccessible\n'
    '{\n'
    '  "status": "degraded",\n'
    '  "checks": {\n'
    '    "database": { "ok": false, "latencyMs": 5003 }\n'
    '  }\n'
    '}'
)
separator()

endpoint_box("GET", "/api/health/detail", "Rapport de santé complet — Admin uniquement")
para("Retourne des métriques détaillées sur tous les composants de l'API :")
code_block(
    '{\n'
    '  "status": "healthy",\n'
    '  "checks": {\n'
    '    "database": {\n'
    '      "ok": true, "latencyMs": 2, "status": "connected"\n'
    '    },\n'
    '    "database_stats": {\n'
    '      "users": 142, "dashboards": 89, "widgets": 512\n'
    '    },\n'
    '    "cache": {\n'
    '      "ok": true, "type": "IMemoryCache"\n'
    '    },\n'
    '    "process": {\n'
    '      "pid": 1234,\n'
    '      "memoryMb": 87.3,\n'
    '      "threadCount": 12,\n'
    '      "uptimeSeconds": 86400\n'
    '    },\n'
    '    "config": {\n'
    '      "jwt": true, "gemini": true, "connStr": true\n'
    '    }\n'
    '  }\n'
    '}'
)

add_table(
    ["Métrique", "Description", "Valeur idéale"],
    [
        ["database.latencyMs",      "Temps de réponse de la base de données",           "< 10 ms"],
        ["process.memoryMb",        "Mémoire RAM utilisée par le processus .NET",        "< 500 MB"],
        ["process.threadCount",     "Nombre de threads actifs",                          "< 50"],
        ["process.uptimeSeconds",   "Temps depuis le dernier démarrage de l'API",        "Le plus grand possible"],
        ["config.gemini",           "La clé API Gemini est-elle configurée ?",           "true"],
        ["config.jwt",              "La clé secrète JWT est-elle configurée ?",          "true"],
    ],
    [4.5, 7.5, 5]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 14. SÉCURITÉ
# ═══════════════════════════════════════════════════════════════════════════════
h1("14. Sécurité Backend")
separator()

h2("14.1 Authentification JWT")
para(
    "JWT (JSON Web Token) est le mécanisme d'authentification central. Après connexion, "
    "l'utilisateur reçoit un token signé contenant ses informations. Ce token est envoyé "
    "dans chaque requête via l'en-tête Authorization: Bearer <token>."
)
bullet("Algorithme de signature : HS256 (HMAC-SHA256)")
bullet("Durée de validité : 7 jours")
bullet("ClockSkew = 0 : les tokens expirés sont rejetés immédiatement (pas de tolérance)")
bullet("Validation : Issuer, Audience, Lifetime, Signature — les 4 paramètres sont vérifiés")

h2("14.2 Hachage des Mots de Passe (BCrypt)")
para(
    "BCrypt est l'algorithme recommandé pour les mots de passe. Contrairement à MD5 ou SHA-1, "
    "BCrypt est intentionnellement lent (cost factor configurable) et inclut un sel automatique. "
    "Même si la base de données est volée, les mots de passe ne peuvent pas être récupérés."
)

h2("14.3 Rate Limiting")
add_table(
    ["Politique", "Limite", "Endpoints concernés", "Objectif"],
    [
        ["auth-policy",   "5 req/min",   "/api/auth/login, /register",       "Anti brute-force"],
        ["ai-policy",     "15 req/min",  "Tous les endpoints /api/ai/*",      "Limiter le coût Gemini"],
        ["upload-policy", "5 req/min",   "/api/datasource/upload-csv",        "Éviter les abus d'upload"],
        ["export-policy", "20 req/min",  "/api/export/*/json, /export/*/pdf", "Éviter la génération massive"],
    ],
    [3, 2.5, 6.5, 5]
)

h2("14.4 Protection contre les Injections SQL")
para(
    "Entity Framework Core génère des requêtes SQL paramétrées automatiquement. "
    "Il est impossible d'injecter du SQL via les paramètres de l'API : EF Core traite "
    "toutes les entrées utilisateur comme des paramètres, jamais comme du code SQL."
)

h2("14.5 Validation des Fichiers CSV")
para(
    "5 niveaux de validation (taille, extension, MIME, magic bytes, contenu) empêchent "
    "l'upload de fichiers malveillants déguisés en CSV. Voir section 6 pour le détail."
)

h2("14.6 En-têtes de Sécurité HTTP")
add_table(
    ["En-tête", "Valeur", "Protection contre"],
    [
        ["X-Content-Type-Options",   "nosniff",         "MIME sniffing attacks"],
        ["X-Frame-Options",          "DENY",            "Clickjacking"],
        ["X-XSS-Protection",         "1; mode=block",   "XSS basique (legacy)"],
        ["Referrer-Policy",          "no-referrer",     "Fuite d'URL dans les requêtes"],
        ["Content-Security-Policy",  "default-src 'none'", "XSS, injection de contenu"],
        ["Permissions-Policy",       "camera=(), mic=()", "Accès aux API navigateur"],
    ],
    [5, 4.5, 7.5]
)

h2("14.7 Gestion Sécurisée des Secrets")
para(
    "Les secrets (clé JWT, clé Gemini, connexion DB) ne sont JAMAIS committés dans Git. "
    "Ils sont placés dans appsettings.Development.json (ajouté dans .gitignore) ou "
    "dans des variables d'environnement. appsettings.json ne contient que des valeurs vides."
)

h2("14.8 Audit Logging")
para(
    "Toutes les actions sensibles sont tracées en base avec :"
)
bullet("UserId : qui a effectué l'action")
bullet("Action : code de l'action (USER_LOGIN, DASHBOARD_DELETED, etc.)")
bullet("EntityType / EntityId : sur quelle entité")
bullet("IpAddress : depuis quelle adresse IP")
bullet("CreatedAt : quand (UTC)")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 15. DOCKER & DÉPLOIEMENT
# ═══════════════════════════════════════════════════════════════════════════════
h1("15. Docker et Déploiement")
separator()

h2("15.1 Pourquoi Docker ?")
para(
    "Docker permet d'empaqueter l'API et toutes ses dépendances dans un conteneur portable. "
    "N'importe qui peut démarrer l'application complète (API + base de données) avec "
    "une seule commande, sur n'importe quel serveur, sans configurer manuellement l'environnement."
)

h2("15.2 Build Multi-Étapes (Dockerfile)")
code_block(
    "# Étape 1 : Compilation du code C#\n"
    "FROM mcr.microsoft.com/dotnet/sdk:8.0 AS build\n"
    "WORKDIR /src\n"
    "COPY DashboardAPI.csproj ./\n"
    "RUN dotnet restore            # Télécharge les packages NuGet\n"
    "COPY . ./\n"
    "RUN dotnet publish -c Release -o /app/publish\n\n"
    "# Étape 2 : Image finale légère (sans le SDK)\n"
    "FROM mcr.microsoft.com/dotnet/aspnet:8.0 AS runtime\n"
    "WORKDIR /app\n"
    "# Utilisateur non-root (sécurité)\n"
    "RUN adduser --system appuser\n"
    "COPY --from=build /app/publish ./\n"
    "USER appuser\n"
    "EXPOSE 8080\n"
    "ENTRYPOINT [\"dotnet\", \"DashboardAPI.dll\"]"
)
para(
    "Le build multi-étapes sépare la compilation (image lourde avec SDK) de l'exécution "
    "(image légère runtime-only). L'image finale fait ~200 MB au lieu de ~800 MB.",
    italic=True
)

h2("15.3 Docker Compose")
code_block(
    "# Démarrer l'application complète :\n"
    "docker-compose up -d\n\n"
    "# L'application tourne sur :\n"
    "# API  : http://localhost:8080\n"
    "# DB   : localhost:3307 (MySQL)\n\n"
    "# Services Docker :\n"
    "# ┌─────────────────────────────────────┐\n"
    "# │  db (MySQL 8.0)                     │\n"
    "# │  - Volume persistant mysql_data     │\n"
    "# │  - Healthcheck toutes les 30s       │\n"
    "# │  - Démarrage ~20-30 secondes        │\n"
    "# └──────────────┬──────────────────────┘\n"
    "#                │ (attend que DB soit healthy)\n"
    "# ┌──────────────▼──────────────────────┐\n"
    "# │  api (DashboardAPI)                 │\n"
    "# │  - Attend que MySQL soit prêt       │\n"
    "# │  - Migrations auto au démarrage     │\n"
    "# │  - Volumes : uploads, logs          │\n"
    "# └─────────────────────────────────────┘"
)

h2("15.4 Variables d'Environnement")
para(
    "Les secrets sont passés via un fichier .env (jamais dans docker-compose.yml directement) :"
)
add_table(
    ["Variable", "Description"],
    [
        ["MYSQL_ROOT_PASSWORD",           "Mot de passe root MySQL"],
        ["MYSQL_PASSWORD",                "Mot de passe de l'utilisateur de l'application"],
        ["JWT_KEY",                        "Clé secrète de signature JWT (min. 32 caractères)"],
        ["GEMINI_API_KEY",                 "Clé API Google Gemini"],
        ["ConnectionStrings__Default",     "Chaîne de connexion MySQL complète"],
    ],
    [6.5, 10.5]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 16. TESTS UNITAIRES
# ═══════════════════════════════════════════════════════════════════════════════
h1("16. Tests Unitaires")
separator()

h2("16.1 Pourquoi les Tests sont Importants")
para(
    "Les tests unitaires vérifient automatiquement que le code fonctionne correctement. "
    "Sans tests, chaque modification du code nécessite de tout tester manuellement. "
    "Avec 32 tests automatisés, on peut détecter immédiatement si une modification casse "
    "une fonctionnalité existante."
)

h2("16.2 Couverture des Tests (32 tests au total)")
add_table(
    ["Fichier de tests", "Nb", "Ce qui est testé"],
    [
        ["WidgetDataServiceTests.cs", "11",
         "Pipeline d'agrégation : COUNT, SUM, AVG, KPI, table, filtres, autorisation, format humain"],
        ["AuditServiceTests.cs",      "5",
         "Persistance des logs, IP par défaut 'unknown', timestamps UTC, champ details"],
        ["AuthServiceTests.cs",       "16",
         "Inscription, connexion, tokens JWT, rôles, validation, cas d'erreur"],
    ],
    [5.5, 1.5, 10]
)

h2("16.3 Exemple de Test (WidgetDataService)")
code_block(
    '[Theory]\n'
    '[InlineData("SUM",   120.0)]\n'
    '[InlineData("AVG",    40.0)]\n'
    '[InlineData("COUNT",   3.0)]\n'
    'public async Task Aggregation_ReturnsCorrectValue(string agg, double expected)\n'
    '{\n'
    '    // Arrange : données de test en mémoire (pas de vraie DB)\n'
    '    var data = new[] {\n'
    '        new { region = "Nord", ventes = 30.0 },\n'
    '        new { region = "Nord", ventes = 50.0 },\n'
    '        new { region = "Nord", ventes = 40.0 }\n'
    '    };\n\n'
    '    // Act : appel du service avec la config\n'
    '    var result = await svc.GetWidgetDataAsync(widgetId, userId);\n\n'
    '    // Assert : vérification du résultat\n'
    '    Assert.Equal(expected, GetKpiValue(result));\n'
    '}'
)

h2("16.4 Infrastructure de Test")
bullet("Framework : xUnit (standard .NET)")
bullet("EF Core InMemory : base de données en mémoire pour les tests (pas de MySQL requis)")
bullet("NullLogger : logger qui ne fait rien (évite les sorties inutiles pendant les tests)")
bullet("Les tests s'exécutent avec : dotnet test (zéro configuration)")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 17. FLOW COMPLET
# ═══════════════════════════════════════════════════════════════════════════════
h1("17. Flow Complet de l'Application")
separator()
para(
    "Voici le parcours complet d'un utilisateur type, du premier accès jusqu'à "
    "l'export de son tableau de bord. Chaque étape correspond à un ou plusieurs "
    "endpoints de l'API."
)

code_block(
    "┌─────────────────────────────────────────────────────┐\n"
    "│  1. INSCRIPTION                                     │\n"
    "│     POST /api/auth/register                         │\n"
    "│     → Compte créé, mot de passe haché BCrypt        │\n"
    "│     → AuditLog USER_REGISTERED                      │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  2. CONNEXION                                       │\n"
    "│     POST /api/auth/login                            │\n"
    "│     → Token JWT reçu (7 jours)                      │\n"
    "│     → AuditLog USER_LOGIN                           │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  3. IMPORT DONNÉES                                  │\n"
    "│     POST /api/datasource/upload-csv                 │\n"
    "│     → 5 validations sécurité                        │\n"
    "│     → CSV parsé, stocké en CachedDataJson           │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  4. RECOMMANDATION IA (optionnel)                   │\n"
    "│     POST /api/ai/recommend-chart                    │\n"
    "│     → Gemini suggère le meilleur graphique          │\n"
    "│     → Réponse mise en cache SHA-256 (1h)            │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  5. CRÉATION DASHBOARD                              │\n"
    "│     POST /api/dashboard                             │\n"
    "│     → AuditLog DASHBOARD_CREATED                    │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  6. CRÉATION WIDGET                                 │\n"
    "│     POST /api/widget                                │\n"
    "│     → Widget lié au dashboard, type + config        │\n"
    "│     → AuditLog WIDGET_CREATED                       │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  7. CALCUL DES DONNÉES                              │\n"
    "│     GET /api/widget/{id}/data                       │\n"
    "│     → Pipeline: Load → Filter → Group → Aggregate   │\n"
    "│     → JSON Chart.js retourné au frontend            │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  8. SAUVEGARDE VERSION                              │\n"
    "│     POST /api/dashboard-version/{id}/save           │\n"
    "│     → Snapshot JSON créée (Ctrl+S historique)       │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  9. PARTAGE                                         │\n"
    "│     POST /api/dashboard-share                       │\n"
    "│     → Token GUID généré, lien de partage créé       │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  10. EXPORT PDF                                     │\n"
    "│      GET /api/export/{id}/pdf                       │\n"
    "│      → PDF généré, AuditLog EXPORT_PDF              │\n"
    "└──────────────────────┬──────────────────────────────┘\n"
    "                       │\n"
    "┌──────────────────────▼──────────────────────────────┐\n"
    "│  ADMIN : SURVEILLANCE                               │\n"
    "│  GET /api/stats/overview  → métriques plateforme    │\n"
    "│  GET /api/stats/audit     → activité utilisateurs   │\n"
    "│  GET /api/health/detail   → état serveur + mémoire  │\n"
    "└─────────────────────────────────────────────────────┘"
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 18. CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
h1("18. Conclusion")
separator()

h2("18.1 Récapitulatif des Points Forts")
add_table(
    ["Axe", "Implémentation", "Impact"],
    [
        ["Sécurité",
         "JWT, BCrypt, rate limiting, magic bytes, CSP headers, audit logs",
         "Protection contre les attaques les plus courantes (OWASP Top 10)"],
        ["Performance",
         "AsNoTracking(), cache 5min (stats), cache 1h (IA), pagination",
         "Temps de réponse réduits, scalabilité améliorée"],
        ["Intelligence",
         "4 endpoints Gemini : recommandation, KPI, analyse, chat",
         "Valeur ajoutée différenciante : l'IA guide l'utilisateur"],
        ["Traçabilité",
         "AuditLog sur toutes les actions CRUD + auth + export",
         "Compliance, sécurité, débogage facilité"],
        ["Résilience",
         "Health checks liveness/readiness, auto-refresh background service",
         "Détection proactive des pannes, récupération automatique"],
        ["Qualité",
         "32 tests unitaires, 0 erreur, architecture en couches",
         "Code maintenable, régressions détectées automatiquement"],
        ["Déploiement",
         "Docker multi-stage, docker-compose, variables d'environnement",
         "Déploiement reproductible sur n'importe quel serveur"],
    ],
    [3, 6, 8]
)

h2("18.2 Architecture Moderne et Professionnelle")
para(
    "Ce backend suit les meilleures pratiques de l'industrie .NET :"
)
bullet("Architecture en couches (Controller → Service → DAO → DB) respectant la séparation des responsabilités")
bullet("Injection de dépendances native ASP.NET Core pour un couplage faible et une testabilité maximale")
bullet("Patterns modernes C# 12 : primary constructors, collection expressions, pattern matching")
bullet("API REST standardisée avec format de réponse uniforme ApiResponse<T> sur tous les endpoints")
bullet("Documentation Swagger automatique générée depuis le code (source of truth unique)")

h2("18.3 Chiffres Clés du Projet")
add_table(
    ["Indicateur", "Valeur"],
    [
        ["Nombre de modules API",           "11 modules"],
        ["Nombre d'endpoints",              "50+ endpoints REST"],
        ["Tests unitaires",                 "32 tests (100% passants)"],
        ["Niveaux de validation sécurité",  "5 (upload CSV)"],
        ["Politiques de rate limiting",     "4 politiques distinctes"],
        ["Durée cache IA",                  "1 heure (SHA-256 keyed)"],
        ["Durée cache stats",               "5 minutes"],
        ["Logs conservés",                  "30 jours (api), 90 jours (erreurs)"],
        ["Taille image Docker",             "~200 MB (build multi-stage)"],
    ],
    [8, 9]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Ce backend constitue une base solide et extensible pour une plateforme SaaS professionnelle, "
    "démontrant une maîtrise complète des technologies modernes du développement web backend."
)
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = BLUE

# ─── Save ──────────────────────────────────────────────────────────────────────
output = r"c:\wamp64\www\DashboardAPI\Documentation_API_PFE.docx"
doc.save(output)
print(f"Document généré : {output}")
